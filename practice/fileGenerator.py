from docx import Document

doc = Document()
doc.add_heading('Spring Security Architecture Cheat Sheet', 0)

# Core Concepts
doc.add_heading('Core Concepts', level=1)
doc.add_paragraph('- Authentication: Verifies "Who are you?" (Identity confirmation).')
doc.add_paragraph('- Authorization: Determines "What can you do?" (Permissions after identity is confirmed).')
doc.add_paragraph('- Exception Handling: Manages errors like unauthorized access.')

# Authentication Flow
doc.add_heading('1. Authentication Flow', level=1)
doc.add_paragraph('- Request Arrival: User tries to access a protected resource.')
doc.add_paragraph('- AuthenticationFilter: Extracts credentials (e.g., username/password). Creates an Authentication token.')
doc.add_paragraph('- AuthenticationManager: Orchestrates the authentication process using AuthenticationProviders.')
doc.add_paragraph('- AuthenticationProvider (e.g., DaoAuthenticationProvider): Validates credentials using UserDetailsService and returns an authenticated object.')
doc.add_paragraph('- SecurityContextHolder: Stores the Authentication object in a ThreadLocal SecurityContext.')

doc.add_paragraph('Key Components:', style='List Bullet')
doc.add_paragraph('Authentication: User\'s identity and authorities.', style='List Bullet 2')
doc.add_paragraph('AuthenticationFilter: Extracts credentials.', style='List Bullet 2')
doc.add_paragraph('AuthenticationManager: Manages authentication.', style='List Bullet 2')
doc.add_paragraph('AuthenticationProvider: Validates credentials.', style='List Bullet 2')
doc.add_paragraph('UserDetailsService: Loads user data.', style='List Bullet 2')
doc.add_paragraph('SecurityContext: Holds the Authentication object.', style='List Bullet 2')
doc.add_paragraph('SecurityContextHolder: Access to SecurityContext.', style='List Bullet 2')

# Authorization Flow
doc.add_heading('2. Authorization Flow', level=1)
doc.add_paragraph('- Authenticated Request: Proceeds after successful authentication.')
doc.add_paragraph('- FilterSecurityInterceptor: Intercepts access to protected resources.')
doc.add_paragraph('- SecurityMetadataSource: Provides access rules (e.g., antMatchers).')
doc.add_paragraph('- AccessDecisionManager: Makes authorization decisions via AccessDecisionVoters.')

doc.add_paragraph('Types of Managers:', style='List Bullet')
doc.add_paragraph('Unanimous: All must grant access.', style='List Bullet 2')
doc.add_paragraph('Consensus: Majority must grant.', style='List Bullet 2')
doc.add_paragraph('Affirmative: At least one grants access.', style='List Bullet 2')

doc.add_paragraph('- AccessDecisionVoter (e.g., RoleVoter): Votes on access permissions.')
doc.add_paragraph('- Result: Either access granted or AccessDeniedException.')

doc.add_paragraph('Key Components:', style='List Bullet')
doc.add_paragraph('FilterSecurityInterceptor: Guards resources.', style='List Bullet 2')
doc.add_paragraph('SecurityMetadataSource: Defines access rules.', style='List Bullet 2')
doc.add_paragraph('AccessDecisionManager: Decides on authorization.', style='List Bullet 2')
doc.add_paragraph('AccessDecisionVoter: Votes on access.', style='List Bullet 2')

# Exception Handling
doc.add_heading('3. Exception Handling', level=1)
doc.add_paragraph('- ExceptionTranslationFilter: Catches security exceptions.')
doc.add_paragraph('- AccessDeniedException: Handled by AccessDeniedHandler (403 Forbidden).')
doc.add_paragraph('- AuthenticationException: Handled by AuthenticationEntryPoint.')
doc.add_paragraph('  - HTTP Basic: Prompts login via WWW-Authenticate header.')
doc.add_paragraph('  - Form Login: Redirects to a login page.')

doc.add_paragraph('- @ControllerAdvice: Handles exceptions outside Spring Security filters.')

doc.add_paragraph('Key Components:', style='List Bullet')
doc.add_paragraph('ExceptionTranslationFilter: Catches exceptions.', style='List Bullet 2')
doc.add_paragraph('AccessDeniedHandler: Handles access denied.', style='List Bullet 2')
doc.add_paragraph('AuthenticationEntryPoint: Handles unauthenticated access.', style='List Bullet 2')

# Security Configuration
doc.add_heading('4. Security Configuration', level=1)
doc.add_paragraph('- WebSecurityConfigurerAdapter: Base class for security setup.')
doc.add_paragraph('- Multiple Filter Chains: Use multiple configurers with @Order annotation.')
doc.add_paragraph('- HttpSecurity: Fluent API to configure patterns, auth methods, and rules.')

# Spring Security Filter Chain
doc.add_heading('5. Spring Security Filter Chain (FilterChainProxy)', level=1)
doc.add_paragraph('- FilterChainProxy: The primary filter of Spring Security.')
doc.add_paragraph('- Holds a list of SecurityFilterChain instances.')
doc.add_paragraph('- For each request, it matches a chain and applies its filters (e.g., UsernamePasswordAuthenticationFilter).')

# Save the document
doc.save("Spring_Security_Cheat_Sheet.docx")
