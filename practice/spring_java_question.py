from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize the document
doc = Document()
doc.add_heading("Interview Questions and Answers", 0)

# Define section headers and Q&A data
qa_data = {
    "✅ SPRING / SPRING BOOT": [
        ("1. Important annotations used while creating REST APIs",
         "In Spring Boot, we use annotations like @RestController to define a class as a REST endpoint. "
         "@RequestMapping or the shortcut annotations like @GetMapping, @PostMapping, etc., are used to map HTTP requests "
         "to specific handler methods. @PathVariable binds a value from the URL path, and @RequestParam is used to capture "
         "query parameters. @RequestBody helps in binding incoming JSON data to a Java object, and @ResponseBody ensures a "
         "Java object is serialized into JSON as a response."),
        ("2. Advantages and disadvantages of Spring Boot",
         "Spring Boot reduces development effort with auto-configuration, embedded servers like Tomcat, and production-ready metrics. "
         "It simplifies dependency management using starter projects and supports rapid prototyping. However, it may add unnecessary "
         "dependencies, increase startup time in large projects, and reduce fine-grained control over configurations. Developers sometimes "
         "need to override defaults manually to meet advanced use cases."),
        # Continue adding all Spring Q&As here...
    ],
    "✅ MICROSERVICES": [
        ("1. How one microservice calls another in a microservices architecture",
         "Microservices typically communicate using REST APIs via RestTemplate or WebClient. In more advanced setups, service discovery tools like Eureka are used to dynamically locate services. "
         "API gateways are placed in front of services to handle routing, logging, and security. For asynchronous communication, tools like Kafka or RabbitMQ are used for event-based messaging."),
        # Continue...
    ],
    "✅ JAVA CORE": [
        ("1. Topics you are most familiar with in Java 8",
         "I am most familiar with Streams API, Lambda expressions, Functional Interfaces, and the use of Optional. These features allow writing cleaner and more functional-style code. "
         "I’ve used them extensively in data processing, filtering, and transforming collections. I also understand the differences between default and static methods in interfaces."),
        # Continue...
    ],
    "✅ REST / SOAP & GENERAL": [
        ("1. Difference between REST API and SOAP",
         "REST is lightweight, uses HTTP methods directly, and supports formats like JSON or XML. SOAP is a protocol with strict rules and supports only XML. REST is stateless and better for web-based services, "
         "while SOAP is preferred in enterprise settings with strict standards and WS-* features like security and transactions."),
        # Continue...
    ],
    "✅ CODING": [
        ("1. Unique numbers from list using Stream API",
         'List<Integer> list = Arrays.asList(1, 2, 2, 3, 4, 4, 5);\n'
         'List<Integer> unique = list.stream()\n'
         '  .collect(Collectors.groupingBy(Function.identity(), Collectors.counting()))\n'
         '  .entrySet().stream()\n'
         '  .filter(e -> e.getValue() == 1)\n'
         '  .map(Map.Entry::getKey)\n'
         '  .collect(Collectors.toList());\n'
         'System.out.println(unique);'),
        ("2. Group strings based on character similarity",
         'List<String> words = Arrays.asList("abc", "bca", "xyz", "cab", "zyx");\n'
         'Map<String, List<String>> grouped = words.stream()\n'
         '  .collect(Collectors.groupingBy(word -> {\n'
         '    char[] chars = word.toCharArray();\n'
         '    Arrays.sort(chars);\n'
         '    return new String(chars);\n'
         '  }));\n'
         'System.out.println(grouped.values());'),
    ]
}

# Formatting each section
for section, questions in qa_data.items():
    doc.add_heading(section, level=1)
    for q, a in questions:
        p_q = doc.add_paragraph(q)
        p_q.runs[0].bold = True
        p_a = doc.add_paragraph(a)
        p_a.paragraph_format.space_after = Pt(10)

# Save the file
output_path = "Interview_QA_Complete.docx"
doc.save(output_path)
print(f"Document saved as {output_path}")
